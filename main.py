import tkinter as tk
from tkinter import messagebox
from docx import Document
import requests
import os
from docx2pdf import convert

#RESUME BUILDER USING TKINTER
def get_github_info(username):                           #function to get github info
    user_url = f"https://api.github.com/users/{username}"
    repos_url = f"https://api.github.com/users/{username}/repos"

    user_data = requests.get(user_url).json()
    repos_data = requests.get(repos_url).json()

    if "message" in user_data and user_data["message"] == "Not Found":  #if user is not found 
        return None

    return {              #Returns a dictionary with basic GitHub profile info and a list of repository names.
        "name": user_data.get("name", ""),
        "bio": user_data.get("bio", ""),
        "github_link": user_data.get("html_url", ""),
        "repos": [repo["name"] for repo in repos_data]
    }

def generate_resume():      #function to generate or take input using tkinter    also used as button
    full_name = name_entry.get()
    email = email_entry.get()
    phone = phone_entry.get()
    linkedin = linkedin_entry.get()
    github = github_entry.get()
    skills = skills_entry.get("1.0", tk.END).strip()
    education = education_entry.get("1.0", tk.END).strip()
    experience = experience_entry.get("1.0", tk.END).strip()

    if not full_name or not email:        #basic validation email name required
        messagebox.showerror("Missing Info", "Name and Email are required.")
        return

    github_info = get_github_info(github)   #calls github function to retirve github data

    doc = Document()     #creates a word doc and adds personal info 
    doc.add_heading(full_name, level=0)
    doc.add_paragraph(f"Email: {email}")
    doc.add_paragraph(f"Phone: {phone}")
    doc.add_paragraph(f"LinkedIn: {linkedin}")
    if github_info:    #adda github url and bio and also lists github repos
        doc.add_paragraph(f"GitHub: {github_info['github_link']}")
        doc.add_paragraph(f"GitHub Bio: {github_info['bio']}")
        doc.add_heading("Projects", level=1)
        for repo in github_info["repos"]:
            doc.add_paragraph(f"- {repo}", style="List Bullet")

    doc.add_heading("Skills", level=1)
    doc.add_paragraph(skills)

    doc.add_heading("Education", level=1)
    doc.add_paragraph(education)

    doc.add_heading("Experience", level=1)
    doc.add_paragraph(experience)

    # Save Word Doc in file exploxer where vs folder of this project is made
    docx_file = "resume.docx"   
    doc.save(docx_file)

    # Convert to PDF
    try:
        convert(docx_file, "resume.pdf")
        messagebox.showinfo("Success", "Resume saved as 'resume.pdf'.")
    except Exception as e:
        messagebox.showwarning("PDF Conversion Failed", f"Resume saved as DOCX but PDF failed.\nError: {e}")

# GUI Setup  creates main application window
app = tk.Tk()
app.title("Resume Builder")
app.geometry("600x700")

# Entries
def add_label_entry(label, height=1):  #Reusable function to add a label and either an Entry (single-line input) or Text (multi-line input). height=1 means it’s a single-line input.
    tk.Label(app, text=label).pack()
    if height == 1:
        entry = tk.Entry(app, width=60)
        entry.pack()
        return entry
    else:
        entry = tk.Text(app, height=height, width=60)
        entry.pack()
        return entry

name_entry = add_label_entry("Full Name")    #all inputs in tkinter gui
email_entry = add_label_entry("Email")
phone_entry = add_label_entry("Phone Number")
linkedin_entry = add_label_entry("LinkedIn URL")
github_entry = add_label_entry("GitHub Username")

skills_entry = add_label_entry("Skills (comma-separated or list)", height=3)    #creates multi line text boxes
education_entry = add_label_entry("Education", height=3)
experience_entry = add_label_entry("Experience", height=3)

tk.Button(app, text="Generate Resume", command=generate_resume, bg="blue", fg="white").pack(pady=20)    #generate buttuon

app.mainloop()
